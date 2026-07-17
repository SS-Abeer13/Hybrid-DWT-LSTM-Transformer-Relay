import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define workspace directory
workspace_dir = r"f:\Downloads\Transformer Thesis"

# Initialize document
doc = Document()

# Page setup (A4, 1-inch margins)
for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.4)  # MIST standard left margin for binding
    section.right_margin = Inches(1.0)

# Helper function to add centered text
def add_centered_para(doc, text, size_pt, bold=False, italic=False, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    return p

# Helper function for left-aligned body text
def add_body_para(doc, text, size_pt=12, bold=False, italic=False, space_after=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    return p

# Helper function for signature lines
def add_sig_block(doc, title, name, details, space_before=24):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run_title = p.add_run(f"{title}\n")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(12)
    run_title.font.bold = True
    
    # 6.5 cm line (approx 2.5 inches)
    run_line = p.add_run("_________________________________________\n")
    run_line.font.name = 'Times New Roman'
    run_line.font.size = Pt(12)
    
    run_details = p.add_run(f"{name}\n{details}")
    run_details.font.name = 'Times New Roman'
    run_details.font.size = Pt(11)

# ==================== PAGE 1: TITLE PAGE (OUTER COVER) ====================
doc.add_paragraph().paragraph_format.space_before = Pt(36) # Vertical spacing
add_centered_para(doc, "HYBRID ADAPTIVE TRANSFORMER DIFFERENTIAL", 18, bold=True, space_after=6)
add_centered_para(doc, "PROTECTION:", 18, bold=True, space_after=6)
add_centered_para(doc, "AN INTELLIGENT DWT-LSTM FRAMEWORK", 18, bold=True, space_after=100)

add_centered_para(doc, "Md Moinul Haque", 14, bold=True, space_after=4)
add_centered_para(doc, "Md Reyaz Uddin", 14, bold=True, space_after=100)

add_centered_para(doc, "B.Sc. MARITIME SCIENCE THESIS", 14, bold=True, space_after=100)

add_centered_para(doc, "FACULTY OF MARITIME BUSINESS STUDIES", 12, bold=True, space_after=4)
add_centered_para(doc, "BANGLADESH MARITIME UNIVERSITY", 14, bold=True, space_after=4)
add_centered_para(doc, "DHAKA, BANGLADESH", 12, bold=True, space_after=36)

add_centered_para(doc, "JANUARY 2026", 12, bold=True, space_after=0)

# Page Break to Page 2
doc.add_page_break()

# ==================== PAGE 2: SECOND TITLE PAGE (INNER COVER) ====================
doc.add_paragraph().paragraph_format.space_before = Pt(24)
add_centered_para(doc, "HYBRID ADAPTIVE TRANSFORMER DIFFERENTIAL PROTECTION:\nAN INTELLIGENT DWT-LSTM FRAMEWORK", 14, bold=True, space_after=54)

add_centered_para(doc, "Md Moinul Haque (ID: 202216206) & Md Reyaz Uddin (ID: 202216211)", 12, bold=True, space_after=54)

add_centered_para(doc, "A Thesis Submitted in Partial Fulfillment of the Requirements for the\nDegree of Bachelor of Science in Port and Shipping Management", 12, italic=True, space_after=72)

add_centered_para(doc, "FACULTY OF MARITIME BUSINESS STUDIES\nBANGLADESH MARITIME UNIVERSITY\nDHAKA, BANGLADESH", 12, bold=True, space_after=72)

add_centered_para(doc, "JANUARY 2026", 12, bold=True, space_after=0)

# Page Break to Page 3
doc.add_page_break()

# ==================== PAGE 3: APPROVAL CERTIFICATE ====================
doc.add_paragraph().paragraph_format.space_before = Pt(12)
add_centered_para(doc, "APPROVAL CERTIFICATE", 16, bold=True, space_after=24)

p_cert = add_body_para(doc, "This is to certify that the thesis entitled ", space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
p_cert.runs[0].font.italic = True

add_centered_para(doc, "“Hybrid Adaptive Transformer Differential Protection:\nAn Intelligent DWT-LSTM Framework”", 13, bold=True, space_after=18)

add_body_para(doc, "submitted by Md Moinul Haque (ID: 202216206) & Md Reyaz Uddin (ID: 202216211) to the Faculty of Maritime Business Studies, Bangladesh Maritime University (BMU), Dhaka, Bangladesh, in partial fulfillment of the requirements for the degree of Bachelor of Science in Port and Shipping Management, has been accepted as satisfactory for the award of the degree and approved by the undersigned examiners.", space_after=18)

# Examiner signatures
add_sig_block(doc, "Supervisor:", "Cdre A N M Didarul Alam, (L), NUP, psc, BN", "Dean, Faculty of Maritime Business Studies, Bangladesh Maritime University", space_before=18)
add_sig_block(doc, "Head of the Department:", "Head of Department Name", "Faculty of Maritime Business Studies, BMU", space_before=14)
add_sig_block(doc, "Internal Examiner:", "Internal Examiner Name", "Associate Professor, Department of EEE, University", space_before=14)
add_sig_block(doc, "External Examiner:", "External Examiner Name", "Professor, Department of EEE, Other University, City, Bangladesh", space_before=14)

p_counter = doc.add_paragraph()
p_counter.paragraph_format.space_before = Pt(18)
run_counter = p_counter.add_run("Counter-signed on behalf of the Board of Advanced Studies.")
run_counter.font.name = 'Times New Roman'
run_counter.font.size = Pt(11)

# Page Break to Page 4
doc.add_page_break()

# ==================== PAGE 4: DECLARATION ====================
doc.add_paragraph().paragraph_format.space_before = Pt(12)
add_centered_para(doc, "DECLARATION", 16, bold=True, space_after=24)

p_decl = add_body_para(doc, "I hereby declare that the thesis entitled ", space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
p_decl.runs[0].font.italic = True

add_centered_para(doc, "“Hybrid Adaptive Transformer Differential Protection:\nAn Intelligent DWT-LSTM Framework”", 13, bold=True, space_after=18)

add_body_para(doc, "submitted to the Faculty of Maritime Business Studies, Bangladesh Maritime University, Dhaka, Bangladesh, in partial fulfillment of the requirements for the degree of Bachelor of Science in Port and Shipping Management, is a record of original work carried out by me under the supervision of Cdre A N M Didarul Alam, (L), NUP, psc, BN, Dean, Faculty of Maritime Business Studies, Bangladesh Maritime University. I hereby affirm that:", space_after=12)

# Checklist items
checklist = [
    "This thesis has not been submitted elsewhere, in whole or in part, for the award of any degree, diploma, or fellowship at this or any other university, institute, or organization.",
    "The work presented herein is entirely the original research and intellectual contribution of the candidate, except where explicit reference has been made to the work of other researchers.",
    "Proper acknowledgment has been made in the text to all other materials used, and all sources of information have been duly cited following the IEEE referencing format.",
    "No part of this thesis has been fabricated, plagiarized, or misrepresented in any manner, and the results presented are accurate to the best of my knowledge and belief.",
    "I have complied with all ethical guidelines and academic integrity standards set forth by the Bangladesh Maritime University."
]

for item in checklist:
    p_item = doc.add_paragraph(style='List Bullet')
    p_item.paragraph_format.space_after = Pt(4)
    p_item.paragraph_format.left_indent = Inches(0.25)
    run_item = p_item.add_run(item)
    run_item.font.name = 'Times New Roman'
    run_item.font.size = Pt(11)

# Signatures at the bottom
p_sig = doc.add_paragraph()
p_sig.paragraph_format.space_before = Pt(36)
p_sig.paragraph_format.space_after = Pt(4)
run_sig = p_sig.add_run("Candidate's Signature: _________________________________________\n\n")
run_sig.font.name = 'Times New Roman'
run_sig.font.size = Pt(11)

run_names = p_sig.add_run("Names: Md Moinul Haque & Md Reyaz Uddin\n")
run_names.font.name = 'Times New Roman'
run_names.font.size = Pt(11)
run_names.font.bold = True

run_ids = p_sig.add_run("Student IDs: 202216206 & 202216211\n\n")
run_ids.font.name = 'Times New Roman'
run_ids.font.size = Pt(11)

run_date = p_sig.add_run("Date: ________________________")
run_date.font.name = 'Times New Roman'
run_date.font.size = Pt(11)

# Save Document
target_path = os.path.join(workspace_dir, "first_4_pages.docx")
doc.save(target_path)
print(f"SUCCESS: Generated first 4 pages document at: {target_path}")
